from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("/Users/tianqi/Documents/CODEX-MAP/docs/传神智谱优化设计")
PRD_OUT = OUT_DIR / "传神智谱-本体开发平台优化PRD-v0.2.docx"
DESIGN_OUT = OUT_DIR / "传神智谱-本体开发平台详细设计-v0.2.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 39, 54)
MUTED = RGBColor(96, 104, 116)
RED = RGBColor(155, 28, 28)
GREEN = RGBColor(26, 115, 84)
LIGHT_BLUE = "EEF5FF"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "ECF8F2"
LIGHT_RED = "FFF1F1"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_para_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def paragraph_border_bottom(paragraph, color="B7C9DD", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def configure_doc(doc: Document, header_text: str):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = header_text
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(header, after=0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("内部设计文档")
    set_run_font(run, size=9, color=MUTED)


def add_para(doc, text="", size=11, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_para_spacing(p, before=16, after=8)
        size, color = 16, BLUE
    elif level == 2:
        set_para_spacing(p, before=12, after=6)
        size, color = 13, BLUE
    else:
        set_para_spacing(p, before=8, after=4)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_para_spacing(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=color)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_para_spacing(p, after=4, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_callout(doc, title, body, fill=LIGHT_GRAY, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    set_para_spacing(p2, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        set_para_spacing(p, after=0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_para_spacing(p, after=0, line=1.12)
            r = p.add_run(text)
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_title_page(doc, title: str, subtitle: str, doc_type: str, meta: list[tuple[str, str]], conclusion: str):
    add_para(doc, doc_type, size=11, color=MUTED, bold=True, after=8)
    p = doc.add_paragraph()
    set_para_spacing(p, after=4)
    r = p.add_run(title)
    set_run_font(r, size=24, color=INK, bold=True)
    add_para(doc, subtitle, size=13, color=MUTED, after=16)
    for label, value in meta:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=11, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=11, color=INK)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)
    add_callout(doc, "设计结论", conclusion, fill=LIGHT_BLUE)


def build_prd():
    doc = Document()
    configure_doc(doc, "传神智谱 · 本体开发平台优化 PRD")
    add_title_page(
        doc,
        "传神智谱：本体开发平台优化 PRD v0.2",
        "从静态本体 CRUD 升级为合同领域语义能力生产系统",
        "产品需求文档",
        [
            ("版本", "v0.2"),
            ("日期", "2026-06-06"),
            ("分支", "feature/zhitu-ontology-optimization"),
            ("业务域", "合同审核与合同风险决策"),
            ("当前边界", "不做合同创建，不做场景编排，先完善本体开发能力"),
        ],
        "智谱的核心价值不是把对象、关系、规则做成表单，而是让业务专家把合同领域的语义结构、计算口径、行动触发和查询能力沉淀成可测试、可发布、可被智脑调用的能力资产。",
    )

    add_heading(doc, "1. 定位升级", 1)
    add_para(doc, "传神智谱从“本体管理工具”升级为“合同领域本体开发平台”。它面向业务、法务、风控和解决方案团队，不要求用户理解 Neo4j、Cypher 或 OWL，而是通过本体、对象、关系、函数、行为、查询能力和验证用例，把复杂逻辑沉到后台。")
    add_bullet(doc, "智谱负责生产能力：定义合同语义结构、计算指标、行动规则、查询接口和版本。")
    add_bullet(doc, "智脑负责运行能力：调用智谱发布的能力，完成推理、决策、风控和执行协同。")
    add_bullet(doc, "场景不作为智谱一期主功能；场景应在智脑模块中编排，智谱只提供可被场景引用的基础能力。")

    doc.add_page_break()
    add_heading(doc, "2. 产品原则", 1)
    add_table(
        doc,
        ["原则", "产品含义", "落地要求"],
        [
            ["以本体为单位", "进入智谱后先选择本体，再查看对象、关系、行为和图谱。", "左侧一级导航保留“智谱”；二级先是本体库，点击本体后进入详情。"],
            ["属性在对象里", "属性是对象定义的一部分，不独立成为一级管理心智。", "对象详情内维护属性、枚举、计算属性和字段口径。"],
            ["规则在行为里", "规则用于决定行为何时触发、如何输出。", "行为详情内维护触发条件、规则、输出模板和人工确认要求。"],
            ["函数沉淀语义", "复杂业务口径通过函数或指标表达。", "函数可绑定对象属性，也可被行为规则和查询能力引用。"],
            ["每个接口可测试", "新增接口必须同步写测试，不允许只做页面。", "API、负例、引用一致性、审计、图同步都要有测试用例。"],
        ],
        [1.15, 2.55, 2.8],
        font_size=9.2,
    )

    doc.add_page_break()
    add_heading(doc, "3. 信息架构", 1)
    add_para(doc, "智谱一级导航保持简单。用户点击智谱后看到“本体库”，本体是管理和发布的单位。")
    add_table(
        doc,
        ["层级", "页面", "说明"],
        [
            ["一级", "智谱", "企业合同语义能力生产入口。"],
            ["二级", "本体库", "展示所有本体，支持新增、复制、停用、导入、导出。"],
            ["本体内", "总览", "查看对象、关系、函数、行为、查询能力、验证通过率和发布状态。"],
            ["本体内", "对象", "维护对象及其属性、计算属性、字段口径和样例。"],
            ["本体内", "关系", "维护对象之间的业务关系、方向、基数和可遍历性。"],
            ["本体内", "函数与指标", "维护内置函数引用、自定义计算函数、指标口径和测试样例。"],
            ["本体内", "行为与规则", "维护行动模板、触发条件、规则、输出和人工确认要求。"],
            ["本体内", "查询能力", "把图查询、指标查询、多跳推理封装成可调用能力。"],
            ["本体内", "本体图谱", "只展示对象和关系，点击对象看属性，点击边看相关规则和行为。"],
            ["本体内", "验证与发布", "运行样例验证，通过后发布版本给智脑。"],
            ["本体内", "版本与审计", "查看变更、差异、回滚和操作日志。"],
        ],
        [0.8, 1.3, 4.4],
        font_size=9.1,
    )

    doc.add_page_break()
    add_heading(doc, "4. 用户角色", 1)
    add_table(
        doc,
        ["角色", "核心任务", "权限"],
        [
            ["平台管理员", "管理本体空间、成员、发布策略和系统配置。", "全部权限。"],
            ["业务建模人", "维护对象、属性、关系、函数和查询能力。", "编辑草稿、运行验证、提交发布。"],
            ["法务专家", "维护行为规则、审核风险口径和已解案例。", "编辑规则、审批发布、查看验证结果。"],
            ["智脑开发者", "调用已发布本体能力，联调推理和执行闭环。", "只读已发布版本、调用查询能力。"],
            ["审计查看者", "查看版本、变更和发布记录。", "只读。"],
        ],
        [1.2, 3.1, 2.2],
        font_size=9.3,
    )

    doc.add_page_break()
    add_heading(doc, "5. MVP 功能范围", 1)
    add_callout(
        doc,
        "一期必须完成",
        "本体库、对象与属性、关系、本体图谱、函数与指标、行为与规则、查询能力、验证与发布、版本与审计。每个管理模块必须支持新增、查看、编辑、停用、复制、导入、导出或说明为什么暂不支持。",
        fill=LIGHT_GREEN,
        title_color=GREEN,
    )
    add_heading(doc, "5.1 本体库", 2)
    add_bullet(doc, "支持创建合同审核本体，复制历史本体，停用本体，导入 YAML/JSON，导出版本包。")
    add_bullet(doc, "本体卡片展示对象数、关系数、函数数、行为数、查询能力数、验证通过率和发布状态。")
    add_bullet(doc, "禁止物理删除已发布本体；只允许停用或废弃版本。")
    add_heading(doc, "5.2 对象与属性", 2)
    add_bullet(doc, "对象列表只展示对象名称、业务编码、状态、属性数量、关联关系数量和更新时间。")
    add_bullet(doc, "对象详情内维护普通属性、枚举属性、计算属性、字段来源、中文口径和样例值。")
    add_bullet(doc, "计算属性必须绑定函数、输入字段、输出类型和测试样例。")
    add_heading(doc, "5.3 关系", 2)
    add_bullet(doc, "关系连接两个对象，支持方向、基数、业务定义、可遍历性和样例。")
    add_bullet(doc, "关系可以绑定影响它的行为规则，但规则入口仍在行为详情中。")
    add_heading(doc, "5.4 函数与指标", 2)
    add_bullet(doc, "提供内置函数：计数、求和、最大值、最小值、最后一个值、时间窗口聚合、同比、环比、阈值比较。")
    add_bullet(doc, "支持自定义函数草稿、输入输出定义、代码模板、AI 辅助生成、测试运行和版本管理。")
    add_bullet(doc, "一期代码执行可先做沙箱模拟和预置函数，自定义代码执行作为受控开关。")
    add_heading(doc, "5.5 行为与规则", 2)
    add_bullet(doc, "行为是可执行或可建议执行的动作模板，例如触发复核、要求补充材料、阻断生效、生成风险发现。")
    add_bullet(doc, "规则在行为内维护，表达“什么条件下触发该行为”。")
    add_bullet(doc, "行为必须声明输入、输出、副作用、是否需要人工确认、调用智脑后的回写要求。")

    add_heading(doc, "5.6 查询能力", 2)
    add_bullet(doc, "把图查询、指标查询和多跳推理封装成能力，供前端和智脑调用。")
    add_bullet(doc, "每个查询能力必须有输入参数、输出结构、样例请求、样例响应、权限和测试用例。")
    add_heading(doc, "5.7 验证与发布", 2)
    add_bullet(doc, "每个本体版本发布前必须运行验证用例，包括结构校验、引用校验、函数测试、行为规则命中测试和查询能力测试。")
    add_bullet(doc, "验证失败时不允许发布；用户可查看失败原因并跳转到对应对象、函数、行为或查询能力。")
    add_heading(doc, "5.8 暂不做", 2)
    add_bullet(doc, "不做合同创建、合同录入、合同台账、合同审批流和合同在线编辑。")
    add_bullet(doc, "不把场景作为智谱一期主功能；场景编排进入智脑设计。")
    add_bullet(doc, "不做开放式低代码工作流市场；一期只做合同领域可控模板。")

    add_heading(doc, "6. 核心用户流程", 1)
    for step in [
        "进入智谱，打开本体库，选择“单合同审核本体”。",
        "进入对象页，维护合同、主体、条款、义务、付款计划、风险事项等对象，并在对象详情内维护属性。",
        "进入关系页，维护合同包含条款、条款约束主体、付款计划对应义务等关系。",
        "进入函数与指标页，配置付款逾期金额、主体历史风险次数、合同剩余履约天数等计算口径。",
        "进入行为与规则页，配置触发复核、要求补充材料、阻断生效、生成风险发现等行为及规则。",
        "进入查询能力页，把常用图查询封装成能力，例如主体风险穿透、条款偏离检查、付款异常路径。",
        "进入验证与发布页，运行样例验证，通过后发布本体版本。",
    ]:
        add_number(doc, step)

    doc.add_page_break()
    add_heading(doc, "7. 验收标准", 1)
    add_table(
        doc,
        ["模块", "可验收结果"],
        [
            ["本体库", "可以新建、复制、停用、导入、导出本体；进入本体后所有内容按本体隔离。"],
            ["对象与属性", "可以在对象详情内完成属性 CRUD；图谱对象节点只显示中文对象名称。"],
            ["关系", "可以创建对象间关系，图谱边展示关系名称，点击边查看相关行为规则。"],
            ["函数与指标", "可以配置至少 5 类内置函数，并运行样例测试。"],
            ["行为与规则", "可以在行为内维护规则，规则可以引用对象、属性、关系和函数结果。"],
            ["查询能力", "可以创建查询能力并运行测试请求，返回结构化结果。"],
            ["验证与发布", "可以运行本体级验证，失败能定位，成功生成发布版本。"],
            ["接口测试", "每新增一个接口，必须有 pytest API 测试、负例测试和至少一个数据一致性断言。"],
        ],
        [1.35, 5.15],
        font_size=9.4,
    )

    add_heading(doc, "8. 成功指标", 1)
    add_bullet(doc, "业务建模效率：合同审核基础本体从空白到可发布不超过 2 小时。")
    add_bullet(doc, "验证覆盖率：发布前核心对象、关系、函数、行为、查询能力均有测试用例。")
    add_bullet(doc, "智脑可用性：智脑无需拼接散乱 Prompt，可直接调用已发布查询能力和行为规则。")
    add_bullet(doc, "维护稳定性：对象或属性变更时，系统能展示受影响函数、行为、查询能力和版本。")

    add_heading(doc, "9. 路线图", 1)
    add_table(
        doc,
        ["阶段", "目标", "交付"],
        [
            ["P0 文档收敛", "明确智谱定位和实施边界。", "PRD、详细设计、接口测试约束。"],
            ["P1 结构调整", "按本体详情重构导航和页面。", "本体库、对象属性内聚、规则内聚到行为。"],
            ["P2 动态本体", "补齐函数、指标、行为触发和查询能力。", "函数管理、查询能力、验证运行器。"],
            ["P3 发布闭环", "让智脑可调用智谱发布版本。", "发布包、能力目录、调用审计。"],
            ["P4 组织学习", "根据智脑反馈优化智谱。", "反馈回流、规则效果分析、案例库。"],
        ],
        [1.1, 2.5, 2.9],
        font_size=9.4,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(PRD_OUT)


def build_design():
    doc = Document()
    configure_doc(doc, "传神智谱 · 本体开发平台详细设计")
    add_title_page(
        doc,
        "传神智谱：本体开发平台详细设计 v0.2",
        "面向可用接口、可测能力、可发布本体版本的工程实施规格",
        "详细设计文档",
        [
            ("版本", "v0.2"),
            ("日期", "2026-06-06"),
            ("分支", "feature/zhitu-ontology-optimization"),
            ("技术栈", "React + Python/FastAPI + PostgreSQL + Neo4j"),
            ("测试原则", "每新增一个接口，同步完成自动化测试和冒烟验证"),
        ],
        "详细设计采用“PG 主库 + Neo4j 投影 + 本体验证运行器”的架构。PG 保存本体资产、版本、审计和测试结果；Neo4j 保存对象、关系、函数、行为、查询能力之间的语义网络；发布前通过自动化验证保证智脑拿到的是可解释、可调用、可回滚的能力包。",
    )

    add_heading(doc, "1. 总体架构", 1)
    add_para(doc, "本次优化不推翻现有实现，而是在当前 OntologyElement、ReferenceEdge、GraphSyncTask 的基础上做产品结构收敛和模型扩展。短期继续使用通用元素表，避免过早拆出大量物理表；中期可按访问量和治理要求拆成对象、函数、行为、查询能力等专表。")
    add_table(
        doc,
        ["层", "职责", "实现建议"],
        [
            ["前端 React", "本体库、本体详情、CRUD 表单、图谱、验证结果。", "按本体详情页组织 Tab；属性嵌入对象；规则嵌入行为。"],
            ["FastAPI", "REST API、校验、版本、审计、导入导出。", "保留现有通用接口，新增 typed endpoint 或 typed payload。"],
            ["PostgreSQL", "事实源、事务、版本、审计、测试结果。", "继续保存 OntologyElement；新增资源类型 function、query_capability、validation_case。"],
            ["Neo4j", "语义关系投影、影响分析、局部图、多跳查询。", "只投影发布或草稿中的有效元素和引用边。"],
            ["验证运行器", "发布前测试对象、关系、函数、行为、查询能力。", "pytest 覆盖接口；业务验证用例入库并可在 UI 运行。"],
        ],
        [1.25, 2.65, 2.6],
        font_size=9.2,
    )

    add_heading(doc, "2. 资源模型", 1)
    add_table(
        doc,
        ["资源", "resource_type", "说明", "关键 payload"],
        [
            ["本体", "space", "当前由 OntologySpace 承载。", "code、name、domain、template_code、status。"],
            ["对象", "object", "合同领域实体类型。", "aliases、examples、display_config。"],
            ["对象属性", "property", "归属于对象，不做独立导航。", "object_code、data_type、enum_values、source、calculation。"],
            ["关系", "relation", "对象之间的业务连接。", "source_code、target_codes、cardinality、direction、traversable。"],
            ["函数与指标", "function", "业务计算口径。", "function_type、inputs、output、formula、test_cases。"],
            ["行为", "action", "动作模板。", "trigger_scope、input_schema、output_schema、requires_human_review。"],
            ["行为规则", "rule", "归属于行为的触发规则。", "action_code、condition、result、severity、references。"],
            ["查询能力", "query_capability", "封装给前端和智脑调用的能力。", "inputs、query_kind、query_body、output_schema、examples。"],
            ["验证用例", "validation_case", "发布前业务测试。", "case_type、input、expected、assertions。"],
        ],
        [1.05, 1.3, 2.0, 2.15],
        font_size=8.8,
    )

    add_heading(doc, "3. 引用关系", 1)
    add_table(
        doc,
        ["边类型", "起点", "终点", "用途"],
        [
            ["HAS_PROPERTY", "object", "property", "对象拥有属性。"],
            ["RELATES_FROM / RELATES_TO", "relation", "object", "关系的起点和终点。"],
            ["COMPUTED_BY", "property", "function", "计算属性由函数产出。"],
            ["FUNCTION_USES", "function", "property/object/relation", "函数引用输入字段或图结构。"],
            ["ACTION_HAS_RULE", "action", "rule", "行为包含触发规则。"],
            ["RULE_REFERENCES", "rule", "object/property/relation/function", "规则引用业务元素。"],
            ["QUERY_REFERENCES", "query_capability", "object/relation/function/action", "查询能力依赖的本体元素。"],
            ["VALIDATES", "validation_case", "function/action/query_capability", "验证用例覆盖的能力。"],
        ],
        [1.55, 1.3, 1.55, 2.1],
        font_size=9.0,
    )

    doc.add_page_break()
    add_heading(doc, "4. API 设计与测试约束", 1)
    add_callout(
        doc,
        "接口开发红线",
        "每写一个接口必须同时提交测试。测试至少包含成功路径、参数校验失败、引用不存在、状态限制、审计或版本断言。涉及图谱的接口还要断言 ReferenceEdge 或 GraphSyncTask。",
        fill=LIGHT_RED,
        title_color=RED,
    )
    add_heading(doc, "4.1 基础 CRUD 接口", 2)
    add_table(
        doc,
        ["接口", "用途", "必须测试"],
        [
            ["GET /api/spaces", "本体库列表。", "返回默认本体；分页或排序稳定。"],
            ["POST /api/spaces", "创建本体。", "成功创建；重复 code 返回 409；写入审计。"],
            ["PUT /api/spaces/{id}", "编辑本体元信息。", "名称和描述更新；不存在返回 404。"],
            ["POST /api/spaces/{id}/deactivate", "停用本体。", "已发布本体不可物理删除；状态变更写版本。"],
            ["GET /api/ontology/{space_id}/objects", "对象列表。", "分页、搜索、状态筛选；不返回 deleted。"],
            ["POST /api/ontology/{space_id}/objects", "创建对象。", "code 唯一；payload schema 校验；创建版本和图任务。"],
            ["PUT /api/ontology/{space_id}/objects/{id}", "编辑对象。", "版本号递增；影响属性和引用边不丢失。"],
            ["POST /api/ontology/{space_id}/objects/{id}/deactivate", "停用对象。", "有引用时返回影响范围或 require_confirm。"],
            ["POST /api/ontology/{space_id}/objects/{id}/properties", "在对象内创建属性。", "自动创建 HAS_PROPERTY；字段中文名和类型必填。"],
            ["PUT /api/ontology/{space_id}/properties/{id}", "编辑对象属性。", "计算属性绑定函数时校验 function 存在。"],
            ["POST /api/ontology/{space_id}/relations", "创建关系。", "source/target 对象存在；RELATES_FROM/TO 写入。"],
        ],
        [2.15, 1.65, 2.7],
        font_size=8.1,
    )
    doc.add_page_break()
    add_heading(doc, "4.2 能力与发布接口", 2)
    add_table(
        doc,
        ["接口", "用途", "必须测试"],
        [
            ["POST /api/ontology/{space_id}/functions", "创建函数或指标。", "输入输出 schema 校验；样例测试可运行。"],
            ["POST /api/ontology/{space_id}/functions/{id}/test", "运行函数测试。", "返回 actual、expected、passed；失败不污染版本。"],
            ["POST /api/ontology/{space_id}/actions", "创建行为。", "输入输出结构合法；默认需要人工确认策略可配置。"],
            ["POST /api/ontology/{space_id}/actions/{id}/rules", "在行为内创建规则。", "规则必须归属行为；引用对象、属性、函数存在。"],
            ["POST /api/ontology/{space_id}/query-capabilities", "创建查询能力。", "输入、输出、查询类型、引用边齐全。"],
            ["POST /api/ontology/{space_id}/query-capabilities/{id}/test", "测试查询能力。", "Neo4j 可用时返回结果；不可用时明确降级。"],
            ["GET /api/ontology/{space_id}/graph", "读取本体图谱。", "只返回对象和关系；对象属性和边上规则通过详情接口返回。"],
            ["POST /api/ontology/{space_id}/validation/run", "运行本体验证。", "失败时阻断发布；结果入库。"],
            ["POST /api/ontology/{space_id}/publish", "发布本体版本。", "未通过验证不可发布；生成发布快照。"],
        ],
        [2.15, 1.65, 2.7],
        font_size=8.1,
    )

    doc.add_page_break()
    add_heading(doc, "5. 测试分层", 1)
    add_table(
        doc,
        ["测试类型", "范围", "执行方式", "通过标准"],
        [
            ["单元测试", "CRUD service、payload schema、引用校验。", "pytest，使用测试数据库事务回滚。", "核心分支覆盖；错误码稳定。"],
            ["API 集成测试", "FastAPI 路由、数据库写入、版本审计。", "TestClient 或 requests。", "每个新增接口至少 3 个用例。"],
            ["图同步测试", "ReferenceEdge、GraphSyncTask、Neo4j 投影。", "Neo4j 可用时集成测试；不可用时验证 outbox。", "图节点、边、状态与主库一致。"],
            ["前端冒烟测试", "本体库、本体详情、表单、图谱、验证运行。", "Playwright。", "核心路径无空白、无报错、按钮可用。"],
            ["发布门禁", "验证用例、函数测试、查询测试。", "发布 API 内部调用验证运行器。", "全部通过才能发布。"],
        ],
        [1.25, 2.0, 1.75, 1.5],
        font_size=9.0,
    )
    add_heading(doc, "5.1 接口开发 DoD", 2)
    for item in [
        "接口有明确请求/响应 schema，不直接透传未校验 payload。",
        "成功路径测试通过，并断言数据库结果。",
        "至少一个负例测试通过，例如重复 code、引用不存在、状态不允许。",
        "涉及增删改时断言 VersionRecord 和 AuditLog。",
        "涉及本体引用时断言 ReferenceEdge；涉及图谱时断言 GraphSyncTask 或 Neo4j 结果。",
        "前端使用该接口的页面完成一次手动或 Playwright 冒烟验证。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 前端页面设计", 1)
    add_table(
        doc,
        ["页面", "关键交互", "实现重点"],
        [
            ["本体库", "卡片/列表切换，新增、复制、停用、导入、导出。", "每个本体显示验证状态和发布状态。"],
            ["本体总览", "指标、待处理问题、快捷入口、最近变更。", "突出是否可发布，不堆概念说明。"],
            ["对象详情", "属性内嵌表格，支持普通属性和计算属性。", "字段尽量中文展示；编码收起到高级区。"],
            ["关系页", "左右对象选择，关系方向和基数配置。", "创建时必须选起点对象和终点对象。"],
            ["函数与指标", "内置函数选择、自定义函数编辑、样例测试。", "先做预置函数，代码函数加风险提示。"],
            ["行为与规则", "行为基础信息 + 规则列表 + 触发测试。", "规则不再做独立主导航。"],
            ["查询能力", "参数定义、测试请求、结构化响应预览。", "面向智脑调用，展示能力说明和样例。"],
            ["本体图谱", "浅色 3D 感对象关系图，节点可拖拽。", "对象节点只显示名称；点击对象看属性，点击边看规则。"],
            ["验证与发布", "一键运行验证，失败跳转定位。", "发布按钮由验证状态控制。"],
        ],
        [1.35, 2.35, 2.8],
        font_size=8.7,
    )

    doc.add_page_break()
    add_heading(doc, "7. 发布包结构", 1)
    add_para(doc, "发布给智脑的不是数据库全量记录，而是版本化能力包。智脑按 capability_id 调用查询能力，按 action_code 理解行为建议，按 rule_id 追溯触发依据。")
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["package_id", "发布包 ID。"],
            ["space_id / version", "来源本体和版本。"],
            ["objects", "对象及属性口径，包含中文显示名和数据类型。"],
            ["relations", "对象关系、方向、基数和可遍历性。"],
            ["functions", "函数名称、输入输出、说明和受控执行方式。"],
            ["actions", "行为模板、规则、输出结构和人工确认策略。"],
            ["query_capabilities", "可调用查询能力列表、参数和样例。"],
            ["validation_report", "发布前验证摘要。"],
        ],
        [1.7, 4.8],
        font_size=9.3,
    )

    add_heading(doc, "8. 与当前代码的落地关系", 1)
    add_table(
        doc,
        ["当前实现", "优化动作"],
        [
            ["OntologySpace", "保留作为本体空间；补 PUT、deactivate、publish。"],
            ["OntologyElement", "继续承载 object/property/relation/action/rule；新增 function、query_capability、validation_case。"],
            ["ReferenceEdge", "扩展边类型，支撑函数、行为规则、查询能力和验证覆盖。"],
            ["GraphProjector", "图谱接口默认只返回对象和关系；详情通过对象/关系接口补充。"],
            ["frontend main.tsx", "把属性和规则从二级导航移入对象/行为详情；新增函数、查询能力、验证发布页面。"],
            ["validation-suite", "升级为接口测试门禁；新增 pytest 或保留 requests smoke，但每个新接口必须覆盖。"],
        ],
        [1.65, 4.85],
        font_size=9.3,
    )

    add_heading(doc, "9. 实施顺序", 1)
    for step in [
        "后端补齐空间更新、停用和发布草案接口，并为现有 CRUD 建立测试基线。",
        "前端调整信息架构：智谱 -> 本体库 -> 本体详情；属性进入对象，规则进入行为。",
        "后端新增 function、query_capability、validation_case 资源类型和引用边类型。",
        "实现函数测试和查询能力测试接口，先支持预置函数和 Neo4j 查询降级。",
        "实现验证与发布：验证失败阻断发布，发布成功生成版本快照。",
        "用 Playwright 验证本体库、对象属性、行为规则、图谱、验证发布五条主路径。",
    ]:
        add_number(doc, step)

    add_heading(doc, "10. 风险与处理", 1)
    add_table(
        doc,
        ["风险", "处理方式"],
        [
            ["过早做成通用低代码平台", "一期限定合同领域模板和受控函数，避免抽象过度。"],
            ["把规则散落到独立表单", "规则必须归属于行为，便于智脑理解动作触发。"],
            ["图谱展示过复杂", "默认只展示对象和关系；属性、规则按点击查看。"],
            ["接口缺测试导致页面能点但不可用", "所有接口按 DoD 做测试，发布前跑全量 API suite。"],
            ["Neo4j 不可用阻塞业务", "PG 为事实源；Neo4j 失败时保留 outbox 并显示同步状态。"],
        ],
        [1.9, 4.6],
        font_size=9.2,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DESIGN_OUT)


def main():
    build_prd()
    build_design()
    print(PRD_OUT)
    print(DESIGN_OUT)


if __name__ == "__main__":
    main()

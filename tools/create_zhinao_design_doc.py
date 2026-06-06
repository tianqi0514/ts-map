from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/tianqi/Documents/CODEX-MAP/docs/传神智脑设计/传神智脑-合同语义网络决策中枢设计.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 39, 54)
MUTED = RGBColor(96, 104, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = RGBColor(168, 32, 40)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph_border_bottom(paragraph, color="B7C9DD", size="8") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_para_spacing(paragraph, before=0, after=6, line=1.1) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_para(doc, text="", size=11, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_para_spacing(p, before=16, after=8)
        size, color = 16, BLUE
    elif level == 2:
        set_para_spacing(p, before=12, after=6)
        size, color = 13, BLUE
    else:
        set_para_spacing(p, before=8, after=4)
        size, color = 12, DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_para_spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_para_spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_para_spacing(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        p = hdr[idx].paragraphs[0]
        set_para_spacing(p, after=0)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_para_spacing(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=10, color=INK)
    set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    doc.add_paragraph()
    return table


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "传神智脑 · 合同语义网络决策中枢"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(header, after=0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("内部设计文档")
    set_run_font(run, size=9, color=MUTED)


def add_title_page(doc: Document) -> None:
    add_para(doc, "产品设计文档", size=11, color=MUTED, bold=True, after=8)
    title = doc.add_paragraph()
    set_para_spacing(title, after=4)
    r = title.add_run("传神智脑：合同语义网络决策中枢设计")
    set_run_font(r, size=24, color=INK, bold=True)
    subtitle = add_para(
        doc,
        "面向组织层面的合同推理、决策、风控与执行协同，而非替代单个员工的执行型 Agent",
        size=13,
        color=MUTED,
        after=16,
    )
    meta = [
        ("适用阶段", "传神智谱 MVP 之后，智脑模块一期设计"),
        ("业务域", "合同全生命周期、合同组合风险、组织级经营决策"),
        ("核心输入", "智谱本体资产、合同事实、组织数据、外部风险数据、历史决策结果"),
        ("核心输出", "风险判断、决策建议、执行编排、推理轨迹、组织记忆"),
        ("日期", "2026-06-06"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=11, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=11, color=INK)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)
    add_callout(
        doc,
        "设计结论",
        "传神智脑不是“员工替身”，而是合同领域的组织级语义网络与决策系统。它把合同、主体、条款、义务、付款、风险、审批、履约和外部事件连接成可推理的业务网络，让企业能从单份合同审核升级到合同组合风险、组织敞口、政策影响和执行闭环管理。",
        fill="EEF5FF",
    )


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)

    add_heading(doc, "1. 模块定位", 1)
    add_para(
        doc,
        "传神智脑是企业合同领域的组织级智能决策中枢。它接收传神智谱发布的本体、关系、行为和规则语义资产，把合同业务事实投影成运行时语义网络，并在这个网络上完成推理、风险识别、决策建议和执行协同。",
    )
    add_para(
        doc,
        "它不以“替代某个法务、采购、财务或销售人员执行任务”为目标。单员工执行型 Agent 的价值边界较窄，容易停留在写邮件、查条款、生成摘要等局部效率提升。智脑的目标是组织层面的合同智能：看清合同背后的主体网络、义务网络、资金网络、审批网络和风险传导网络。",
    )
    add_table(
        doc,
        ["不是", "而是"],
        [
            ["不是单份合同的问答助手", "合同事实、组织关系和外部风险之间的语义网络"],
            ["不是替代员工点击系统按钮", "识别风险、提出决策、生成行动方案、追踪闭环"],
            ["不是把规则硬编码成流程", "把本体、规则、案例和推理轨迹组织成可解释决策"],
            ["不是一次性审核工具", "持续学习的合同经营与风控记忆系统"],
        ],
        [3.05, 3.45],
    )

    add_heading(doc, "2. 与传神智谱的关系", 1)
    add_para(doc, "智谱是定义层，智脑是运行层。智谱沉淀“业务世界如何被理解”，智脑负责“基于这些定义在真实业务中如何判断和行动”。")
    add_table(
        doc,
        ["层级", "职责", "例子"],
        [
            ["传神智谱", "管理本体资产：对象、关系、属性、行为、规则、版本。", "合同、签约主体、条款、付款计划、资质、风险项；对象之间的关系；行为模板。"],
            ["传神智脑", "使用智谱资产构建运行时语义网络，并做推理、决策、风控和执行协同。", "判断某供应商合同组合是否形成集中风险；政策变化影响哪些合同；付款义务是否与履约进度冲突。"],
            ["业务系统", "提供事实和执行回写。", "CLM、ERP、OA、CRM、财务系统、电子签、外部工商/司法/舆情数据。"],
        ],
        [1.25, 2.6, 2.65],
    )
    add_callout(
        doc,
        "边界原则",
        "智谱不保存具体业务单据的全生命周期状态；智脑可以消费合同事实和运行时事件，但一期不替代 CLM、ERP、OA 等系统成为业务主库。",
    )

    add_heading(doc, "3. 核心用户与组织价值", 1)
    add_table(
        doc,
        ["角色", "他们关心什么", "智脑提供什么"],
        [
            ["经营管理者", "合同带来的收入、成本、风险敞口和经营承诺是否可控。", "组织级合同风险地图、供应商/客户集中度、重大合同预警、经营影响分析。"],
            ["法务负责人", "规则是否一致执行，例外是否可追溯，政策变更影响哪些合同。", "规则命中解释、合同组合规则覆盖、政策变更影响面、复盘闭环。"],
            ["采购/销售负责人", "交易能不能推进，卡点在哪里，风险如何处理。", "主体风险穿透、条款谈判建议、执行行动清单、跨部门协同状态。"],
            ["财务/风控", "付款、发票、履约和对手方风险是否联动异常。", "资金义务网络、付款条件与交付里程碑校验、对手方信用与合同金额联动。"],
            ["Agent/系统团队", "如何稳定调用语义能力，不靠临时 Prompt 拼接。", "语义包调用、推理 API、证据链、评测反馈和运行监控。"],
        ],
        [1.25, 2.25, 3.0],
    )

    add_heading(doc, "4. 合同语义网络的运行时对象", 1)
    add_para(doc, "智脑需要在智谱定义的本体之上增加运行时对象。这些对象不是本体定义本身，而是真实业务事实、推理过程和执行闭环。")
    add_table(
        doc,
        ["运行时对象", "含义", "主要关系"],
        [
            ["ContractFact", "从合同文本、台账或业务系统抽取出的合同事实。", "实例化 Contract、Clause、PaymentTerm、Obligation 等本体对象。"],
            ["SemanticCase", "一次组织级合同判断任务。", "关联输入事实、激活场景、命中规则、推理轨迹和决策结论。"],
            ["RiskFinding", "可解释风险发现。", "指向合同、主体、条款、义务、外部证据和触发规则。"],
            ["DecisionProposal", "智脑给出的决策建议。", "包含通过、驳回、补充材料、升级审批、条款修改、执行动作。"],
            ["ActionPlan", "跨系统执行计划。", "分配责任部门、动作、截止时间、依赖、状态和回写结果。"],
            ["Evidence", "支撑判断的证据。", "来自合同原文、组织数据、外部数据、历史案例、规则定义。"],
            ["ReasoningTrace", "推理轨迹。", "记录事实抽取、图遍历、规则命中、冲突处理和最终判断。"],
        ],
        [1.45, 2.45, 2.6],
    )

    add_heading(doc, "5. 功能模块设计", 1)
    add_heading(doc, "5.1 合同语义网络", 2)
    add_para(doc, "负责把合同事实、组织主数据、外部风险数据和智谱本体投影成运行时图网络。用户不需要看到图数据库概念，但可以看到合同背后的组织关系、义务链路、资金链路和风险传导。")
    for item in [
        "单合同网络：从一份合同出发，查看主体、条款、付款、义务、审批、附件和证据。",
        "主体网络：从签约主体出发，穿透股东、实际控制人、历史合作、黑名单和关联方。",
        "义务网络：把交付、付款、验收、发票、违约责任连接起来，识别履约风险。",
        "组合网络：按客户、供应商、事业部、合同类型、金额区间聚合合同敞口。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.2 场景推理舱", 2)
    add_para(doc, "场景不是简单筛选条件，而是组织判断任务的上下文。智脑根据合同事实和用户目标激活相应语义包，决定需要哪些对象、关系、规则、证据和行动。")
    add_table(
        doc,
        ["场景", "典型问题", "输出"],
        [
            ["重大合同决策", "是否可以签，需不需要升级审批，哪些条款必须修改？", "决策建议、风险等级、审批路径、条款修改建议。"],
            ["主体风险穿透", "对手方背后的控制关系和外部风险是否会影响交易？", "主体风险图、关联风险、证据链、处置建议。"],
            ["履约与付款风控", "付款计划、交付里程碑、验收和发票是否一致？", "异常点、资金风险、需补材料、执行动作。"],
            ["政策变更影响", "新政策、新监管或内部制度变化影响哪些存量合同？", "影响合同清单、风险分层、整改计划。"],
            ["合同组合经营分析", "某客户/供应商/业务线的合同敞口和风险趋势如何？", "组合看板、集中度、潜在损失、经营建议。"],
        ],
        [1.4, 2.55, 2.55],
    )

    add_heading(doc, "5.3 决策与风控引擎", 2)
    add_para(doc, "该模块不等同于传统规则引擎。它综合硬规则、软规则、图关系、历史案例和证据置信度，形成可解释的组织决策。")
    for item in [
        "硬规则：黑名单、缺强制条款、资质过期、授权无效等必须拦截的条件。",
        "软规则：实缴资本偏低、责任上限过低、自动续约陷阱、历史履约异常等需综合判断的风险。",
        "图推理：通过股权、控制、历史合作、合同义务、付款链路做多跳穿透。",
        "冲突处理：当规则、历史经验和业务例外冲突时，要求给出冲突说明和人工确认点。",
        "置信度：区分事实确定性、抽取置信度、规则适配度和外部数据时效。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.4 执行协同与回写", 2)
    add_para(doc, "智脑不替代系统执行，但要把决策转成可追踪行动。动作可以推送到 OA、CLM、企业微信、飞书、ERP 或人工任务池。")
    add_table(
        doc,
        ["动作类型", "示例", "回写要求"],
        [
            ["补充材料", "要求对手方补充资质证照、授权委托书、银行信息。", "材料状态、上传时间、验证结果。"],
            ["条款修改", "建议加入责任上限、数据合规、保密期限或退出权条款。", "是否采纳、修改版本、责任人。"],
            ["升级审批", "触发重大合同、关联交易、超预算或高风险主体审批。", "审批节点、结论、驳回原因。"],
            ["风险处置", "标记冻结付款、暂停签署、发起尽调或要求担保。", "处置状态、完成证据、残余风险。"],
        ],
        [1.35, 3.05, 2.1],
    )

    add_heading(doc, "5.5 组织记忆与复盘", 2)
    add_para(doc, "每一次智脑判断都应进入组织记忆。系统记录哪些规则被命中、哪些证据被采用、人工是否接受建议、最终业务结果如何，从而支持规则优化和案例沉淀。")
    for item in [
        "规则效果复盘：哪些规则经常被人工推翻，哪些风险经常漏检。",
        "案例沉淀：把典型合同风险、处置过程和最终结果沉淀为可检索案例。",
        "本体反馈：发现新对象、新关系、新属性或新规则需求时，回流到智谱待处理队列。",
        "经营学习：分析不同业务线、区域、客户、供应商的风险模式变化。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 核心页面与体验", 1)
    add_table(
        doc,
        ["页面", "定位", "关键内容"],
        [
            ["智脑工作台", "组织级合同风险与决策入口。", "待决策事项、重大风险、合同敞口、执行卡点、今日需处理。"],
            ["合同语义网络", "从合同或主体出发查看业务网络。", "对象节点、关系边、证据、风险路径、可解释多跳关系。"],
            ["场景推理舱", "选择或自动识别业务场景并运行推理。", "输入事实、激活语义包、命中规则、推理结论、置信度。"],
            ["风险与决策台", "统一管理风险发现和决策建议。", "风险分级、处置建议、责任部门、状态、时限。"],
            ["执行协同台", "把建议变成跨系统动作。", "任务、依赖、回写、提醒、闭环证据。"],
            ["组织记忆", "复盘和持续优化。", "历史案例、人工反馈、规则效果、知识回流。"],
        ],
        [1.35, 2.2, 2.95],
    )

    add_heading(doc, "7. 推理链路", 1)
    add_number(doc, "事实接入：合同文本、台账、审批、付款、发票、履约、主体和外部风险数据进入智脑。")
    add_number(doc, "语义对齐：根据智谱本体把事实映射为对象、属性、关系和证据。")
    add_number(doc, "场景识别：识别当前任务属于重大合同、主体尽调、付款风控、政策影响等哪类场景。")
    add_number(doc, "网络扩展：围绕合同、主体、义务、付款和风险做多跳图遍历，找到相关事实。")
    add_number(doc, "规则与案例推理：执行硬规则检查，结合软规则、历史案例和证据置信度形成判断。")
    add_number(doc, "决策生成：输出风险等级、推荐动作、审批路径和需要人工确认的问题。")
    add_number(doc, "执行闭环：将动作推送到业务系统或责任人，并把结果回写为组织记忆。")

    doc.add_page_break()
    add_heading(doc, "8. 一期 MVP 边界", 1)
    add_callout(
        doc,
        "一期建议",
        "智脑一期不做大而全 Agent 平台，也不做合同系统替代品。建议选择 2 到 3 个高价值组织场景，把语义网络、推理解释和执行闭环跑通。",
        fill="FFF4E5",
    )
    add_table(
        doc,
        ["范围", "一期做", "一期不做"],
        [
            ["业务场景", "主体风险穿透、重大合同风险判断、履约/付款一致性检查。", "全量合同生命周期管理、复杂合同起草协同。"],
            ["数据接入", "使用结构化样例、导入数据、少量合同事实抽取。", "打通所有企业系统、实时全量数据湖。"],
            ["推理能力", "基于本体、规则、图遍历和证据链的可解释推理。", "开放式通用聊天、完全自治决策。"],
            ["执行", "生成行动建议和任务清单，可人工确认后执行。", "自动代替员工完成不可逆业务操作。"],
            ["学习闭环", "记录人工采纳/驳回和原因，回流智谱优化。", "自动改写本体并直接发布。"],
        ],
        [1.25, 2.75, 2.5],
    )

    add_heading(doc, "9. 与现有代码的落地关系", 1)
    add_para(doc, "当前代码已经具备智脑一期所需的一部分底座：本体 CRUD、YAML 导入、Neo4j 图投影、图遍历、验证套件。下一步不是另起炉灶，而是在这些能力上增加运行时事实、推理任务和执行闭环。")
    add_table(
        doc,
        ["已有能力", "可复用方式", "需要补充"],
        [
            ["智谱本体管理", "作为智脑语义包来源。", "语义包发布状态、版本锁定、场景激活范围。"],
            ["ReferenceEdge + Neo4j", "支撑运行时图遍历和影响分析。", "业务事实图、证据图、推理轨迹图。"],
            ["YAML 场景模板", "作为行业本体和验证样例。", "导入字段结构化保留、场景与规则引用关系。"],
            ["GraphTraversalPanel", "作为开发/调试能力。", "业务化的“风险路径”和“证据路径”展示。"],
            ["Validation Suite", "作为智脑推理质量的验收基础。", "加入决策准确率、证据完整性、人工采纳率指标。"],
        ],
        [1.5, 2.45, 2.55],
    )

    add_heading(doc, "10. 数据模型增量建议", 1)
    add_table(
        doc,
        ["模型", "用途", "核心字段"],
        [
            ["RuntimeFact", "保存合同运行时事实。", "source_type、source_id、ontology_code、value、confidence、evidence_id。"],
            ["SemanticCase", "一次推理任务。", "case_type、contract_id、space_id、semantic_package_version、status。"],
            ["DecisionResult", "保存决策输出。", "risk_level、decision、recommended_actions、requires_human_review。"],
            ["ReasoningTrace", "推理过程可追溯。", "step_type、input_refs、rule_refs、evidence_refs、output、confidence。"],
            ["ExecutionTask", "执行协同。", "owner、action_code、due_at、status、external_system、callback_payload。"],
            ["FeedbackRecord", "人工反馈和复盘。", "accepted、override_reason、final_outcome、ontology_feedback。"],
        ],
        [1.45, 2.05, 3.0],
    )

    doc.add_page_break()
    add_heading(doc, "11. 路线图", 1)
    add_table(
        doc,
        ["阶段", "目标", "关键交付"],
        [
            ["P0 概念收敛", "明确智脑不是个人 Agent，而是合同语义网络决策中枢。", "定位文档、边界说明、核心场景清单。"],
            ["P1 语义网络 MVP", "围绕合同审核构建运行时事实图。", "合同语义网络页、主体风险穿透、风险路径展示。"],
            ["P2 推理与决策", "让系统能生成可解释风险判断和决策建议。", "SemanticCase、ReasoningTrace、DecisionResult、规则命中解释。"],
            ["P3 执行闭环", "把建议变成跨部门动作并回写结果。", "ActionPlan、执行任务、人工确认、业务系统回调。"],
            ["P4 组织学习", "形成可持续优化的合同智能能力。", "反馈学习、规则效果分析、案例库、智谱回流机制。"],
        ],
        [1.25, 2.45, 2.8],
    )

    add_heading(doc, "12. 成功指标", 1)
    for item in [
        "风险发现率：高风险合同、主体穿透风险、付款/履约异常是否被稳定发现。",
        "解释完整度：每个结论是否能追溯到事实、证据、规则和图路径。",
        "人工采纳率：法务、风控、业务对决策建议的采纳比例。",
        "执行闭环率：建议动作是否完成、超期、驳回或转人工。",
        "组织学习效率：新风险是否能回流为智谱规则、本体关系或案例。",
        "经营影响：重大合同决策效率、风险损失避免、付款异常减少、审批返工减少。",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "13. 当前必须避免的误区", 1)
    add_table(
        doc,
        ["误区", "为什么危险", "正确做法"],
        [
            ["把智脑做成聊天助手", "会回到个人效率工具，组织级价值不明显。", "围绕合同网络、决策案件和执行闭环组织体验。"],
            ["把规则做成硬编码流程", "合同风险很多是组合判断和证据推理，硬流程会僵化。", "硬规则拦截 + 软规则推理 + 人工确认。"],
            ["让智脑替代 CLM/ERP", "会引入主数据和流程治理复杂度。", "智脑做语义推理和决策协同，业务系统仍是事实源和执行系统。"],
            ["默认全自动执行", "合同决策高风险，容易造成不可逆后果。", "高影响动作必须人工确认，保留证据和审批链。"],
            ["把场景过早做成配置市场", "一期会分散主线，难以验证闭环价值。", "先做合同领域 2 到 3 个深场景。"],
        ],
        [1.45, 2.25, 2.8],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录：推荐的一期页面信息架构", 1)
    add_table(
        doc,
        ["一级模块", "二级页面", "说明"],
        [
            ["智脑", "工作台", "组织级合同风险、待决策事项和执行卡点。"],
            ["智脑", "合同语义网络", "围绕合同、主体、义务、付款、审批展示运行时网络。"],
            ["智脑", "场景推理舱", "选择场景、输入事实、运行推理、查看解释。"],
            ["智脑", "风险与决策", "管理风险发现、决策建议、人工确认和例外。"],
            ["智脑", "执行协同", "将决策建议转成跨部门行动和系统回写。"],
            ["智脑", "组织记忆", "历史案例、反馈、规则效果、本体优化建议。"],
        ],
        [1.2, 1.7, 3.6],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)

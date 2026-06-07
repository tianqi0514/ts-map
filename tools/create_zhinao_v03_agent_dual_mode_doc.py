from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.append(str(Path(__file__).parent))
from create_zhinao_design_doc import (  # noqa: E402
    INK,
    LIGHT_BLUE,
    MUTED,
    add_bullet,
    add_callout,
    add_heading,
    add_number,
    add_para,
    add_table,
    configure_doc,
    paragraph_border_bottom,
    set_para_spacing,
    set_run_font,
)
from create_zhinao_v02_design_doc import (  # noqa: E402
    section_api,
    section_database,
    section_mvp_plan,
    section_product_definition,
    section_reasoning,
    section_risks,
    section_runtime_model,
    section_scenarios,
    section_user_experience,
)


OUT = Path("/Users/tianqi/Documents/CODEX-MAP/docs/传神智脑设计/传神智脑-Agent双模式PRD与详细设计-v0.3.docx")


def add_title_page(doc: Document) -> None:
    add_para(doc, "PRD 与详细设计 · v0.3", size=11, color=MUTED, bold=True, after=8)
    title = doc.add_paragraph()
    set_para_spacing(title, after=4)
    run = title.add_run("传神智脑：Agent 双模式决策平台")
    set_run_font(run, size=24, color=INK, bold=True)
    subtitle = add_para(
        doc,
        "工作流模式负责可控编排，自主 Agent 模式负责意图驱动工具调用；两者都必须和智谱本体发布包建立可追溯链接",
        size=13,
        color=MUTED,
        after=16,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for label, value in [
        ("本次调整", "在 v0.2 的合同语义网络决策平台基础上，补充工作流模式与 Agent 模式的双引擎设计"),
        ("复用策略", "参考 Dify 的节点体系与画布交互；运行时优先采用 LangGraph 思路和 Python 集成；不建议直接 fork 大型平台做重度二开"),
        ("核心约束", "每个节点、工具、技能、推理结果都要支持绑定智谱对象、关系、行为、规则或发布包版本"),
        ("适用场景", "合同风控、重大合同决策、主体风险穿透、履约付款协同、政策影响分析"),
        ("版本日期", "2026-06-07"),
    ]:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=11, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=11, color=INK)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)
    add_callout(
        doc,
        "设计结论",
        "智脑需要同时支持“流程可控”和“意图自主”。工作流适合重大合同审核、付款冻结、补证、审批等高确定性场景；Agent 模式适合用户用自然语言提出复杂问题，由系统自主选择本体查询、图谱遍历、规则执行、技能调用和执行建议。",
        fill="EEF5FF",
    )


def section_open_source_choice(doc: Document) -> None:
    add_heading(doc, "2A. 开源复用选择", 1)
    add_para(
        doc,
        "调研后建议采用“参考 Dify 节点体系 + 后端运行时采用 LangGraph 风格 + 自研本体链接层”的方案。原因是我们的核心差异在合同本体、语义网络和组织级决策，不适合把产品主体直接变成 Dify 或 Flowise 的二开。",
    )
    add_table(
        doc,
        ["方案", "可复用点", "风险/限制", "建议"],
        [
            ["Dify", "节点清单完整，包含 User Input、LLM、Knowledge Retrieval、Agent、If-Else、Human Input、Iteration、Loop、Code、Template、HTTP Request、Tool 等；适合参考画布交互和节点配置。", "许可证是 modified Apache 2.0，前端标识和多租户使用有附加条件；直接 fork web 做商用需法务确认。", "参考节点体系和 DSL 思路，不直接复制前端；必要时仅借鉴后端节点抽象。"],
            ["LangGraph", "Python 生态，适合构建状态图、节点、边、条件路由、工具调用、人机协同和长运行任务。", "它不是低代码 UI，需要我们自建画布和节点配置。", "作为后端运行时首选，和当前 FastAPI/Python 技术栈匹配。"],
            ["Flowise", "视觉化节点编辑成熟，社区活跃，适合参考组件布局。", "Node/TS 后端和现有 Python 后端不一致；近期有安全事件，直接引入要非常谨慎。", "仅参考 UI 交互，不作为核心运行时。"],
            ["完全自研", "最贴合本体、合同风控和组织决策。", "开发周期长，节点执行、调试、版本、运行日志都要自己补齐。", "画布和本体链接层自研，运行时复用 LangGraph，节点体系参考 Dify。"],
        ],
        [1.05, 2.15, 1.85, 1.45],
    )
    add_callout(
        doc,
        "落地判断",
        "一期不要追求完整复制 Dify。我们需要的是：可视化编排体验、标准节点类型、可测试运行时、和智谱本体强绑定。先做最小可用工作流引擎，再逐步扩展节点市场。",
        fill="FFF4E5",
    )


def section_dual_mode_architecture(doc: Document) -> None:
    add_heading(doc, "3A. Agent 双模式架构", 1)
    add_table(
        doc,
        ["模式", "用户心智", "适用任务", "控制方式", "输出"],
        [
            ["工作流模式", "像搭建合同风控流程，明确每一步做什么。", "重大合同审核、固定审批链、付款冻结、补证、政策影响批处理。", "用户自定义节点和边，系统按 DAG/状态图执行。", "每个节点的输入、输出、状态、错误和本体引用。"],
            ["Agent 模式", "像向合同决策专家提问，由系统自己选择工具。", "开放式风险分析、主体穿透、复杂问答、多轮追问、探索式推理。", "Agent 根据意图选择本体工具、技能、图谱查询、业务系统工具。", "最终答复、工具调用日志、推理轨迹、风险与行动建议。"],
            ["混合模式", "流程中某个节点交给 Agent 自主完成。", "固定审批流程中的动态尽调、证据补齐、异常解释。", "工作流约束边界，Agent 在节点内部自主调用工具。", "Agent 子任务结果写回工作流状态。"],
        ],
        [1.1, 1.45, 1.75, 1.85, 1.45],
    )
    add_heading(doc, "3A.1 统一运行状态", 2)
    add_para(doc, "两种模式底层都应共享统一的运行态结构，避免工作流和 Agent 各做一套日志、权限、证据和本体引用。")
    add_table(
        doc,
        ["状态字段", "说明"],
        [
            ["run_id / case_id", "一次工作流或 Agent 推理运行的唯一标识，关联语义案件。"],
            ["package_id / package_version", "锁定本次运行使用的智谱发布包，保证历史可复现。"],
            ["context_facts", "合同事实、主体事实、付款事实、附件事实等运行态事实。"],
            ["ontology_refs", "运行涉及的对象、关系、行为、规则、函数、查询能力引用。"],
            ["evidence_refs", "合同文本、系统记录、外部数据、历史案例等证据。"],
            ["trace_steps", "节点执行、工具调用、规则命中、人工确认、错误重试等轨迹。"],
            ["decision_outputs", "风险发现、决策建议、执行动作和回写结果。"],
        ],
        [1.7, 4.8],
    )


def section_workflow_mode(doc: Document) -> None:
    add_heading(doc, "4A. 工作流模式设计", 1)
    add_para(doc, "工作流模式提供可视化画布，用户可以自定义每个节点、节点之间的连接、运行条件和异常处理。所有节点都支持绑定智谱本体资产。")
    add_heading(doc, "4A.1 节点类型", 2)
    add_table(
        doc,
        ["节点类别", "节点", "用途", "本体链接"],
        [
            ["输入", "用户输入、触发器、文件上传、业务系统事件", "收集合约编号、合同文本、主体编号、审批事件等。", "绑定场景、对象类型、输入字段与属性。"],
            ["语义", "本体包加载、对象映射、属性抽取、关系构建", "把输入事实映射到智谱对象、属性和关系。", "必须选择 package、object、property、relation。"],
            ["LLM", "LLM 生成、文档抽取、问题分类、模板转换", "摘要、条款抽取、风险解释、文案生成。", "提示词可引用对象/规则定义。"],
            ["知识/图谱", "知识检索、图谱查询、路径扩展、影响面分析", "检索案例、读取本体图、运行时图、多跳路径。", "绑定查询能力和关系类型。"],
            ["规则/函数", "规则执行、函数计算、阈值判断、变量聚合", "执行合同金额、主体信用、附件核验等规则。", "绑定行为、规则、函数。"],
            ["控制", "If-Else、循环、迭代、并行、等待、错误处理", "按风险等级、证据完整性、人工意见进行路由。", "条件可引用规则结果和对象属性。"],
            ["人工", "人工输入、人工审批、补证确认、决策确认", "高风险动作进入人工确认。", "绑定行为动作和审批责任角色。"],
            ["工具", "HTTP 请求、代码、业务系统连接器、MCP 工具", "调用 CLM、ERP、OA、企微、外部风控接口。", "工具参数可从对象属性和证据中生成。"],
            ["输出", "答案、风险报告、执行任务、回写、发布", "生成报告、任务、回写业务系统。", "输出绑定决策、行动和组织记忆。"],
        ],
        [0.95, 1.45, 2.15, 1.95],
    )
    add_heading(doc, "4A.2 节点通用配置", 2)
    for item in [
        "基础信息：节点名称、类型、描述、启用状态、超时时间、重试次数。",
        "输入输出：输入变量、输出变量、JSON Schema、字段映射、默认值。",
        "本体链接：space_id、package_id、ontology_refs，支持对象、属性、关系、行为、规则、函数、查询能力。",
        "权限与安全：是否允许外部调用、是否需要人工确认、是否可访问敏感字段。",
        "观测与测试：节点级测试数据、运行日志、耗时、错误、输入输出快照。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4A.3 合同风控工作流示例", 2)
    add_number(doc, "触发器：合同提交审核，读取 contract_id。")
    add_number(doc, "本体包加载：加载合同风控决策本体发布包。")
    add_number(doc, "事实映射：将合同金额、主体、条款、付款计划、附件映射到智谱对象属性。")
    add_number(doc, "图谱查询：扩展合同主体、签署授权、历史风险和关联方路径。")
    add_number(doc, "规则执行：执行高金额、低信用主体、缺失授权、附件未核验规则。")
    add_number(doc, "条件分支：阻断风险进入阻断签署；高风险进入法务/财务审批；中风险进入补证。")
    add_number(doc, "人工审批：责任人确认或覆盖智脑建议。")
    add_number(doc, "执行回写：生成执行任务，回写 CLM/OA 状态。")
    add_number(doc, "组织记忆：记录采纳结果，生成智谱优化建议。")


def section_agent_mode(doc: Document) -> None:
    add_heading(doc, "5A. Agent 模式设计", 1)
    add_para(doc, "Agent 模式把本体、规则、函数、查询能力、技能和业务系统封装成工具。用户只提出目标，Agent 自主选择工具组合，但必须受到权限、成本、迭代次数和人工确认规则约束。")
    add_heading(doc, "5A.1 工具注册表", 2)
    add_table(
        doc,
        ["工具类型", "工具示例", "说明"],
        [
            ["本体工具", "ontology.get_schema、ontology.search、ontology.get_rules、ontology.get_actions", "读取智谱对象、关系、行为、规则、函数定义。"],
            ["图谱工具", "graph.query、graph.expand_paths、graph.impact_analysis", "查询本体图和运行时语义图，返回路径和证据。"],
            ["规则工具", "rule.run、rule.explain、function.evaluate", "执行规则和函数，返回命中结果与解释。"],
            ["事实工具", "fact.extract、fact.map_to_ontology、evidence.retrieve", "抽取合同事实、映射本体、收集证据。"],
            ["技能工具", "skill.contract_summary、skill.party_due_diligence、skill.payment_check", "可复用合同分析能力。"],
            ["业务工具", "clm.get_contract、oa.create_approval、erp.freeze_payment、msg.notify", "连接外部系统，默认高影响动作需人工确认。"],
            ["记忆工具", "memory.search_cases、memory.write_feedback、ontology_feedback.create", "读取历史案例，写入反馈和智谱优化建议。"],
        ],
        [1.15, 2.15, 3.2],
    )
    add_heading(doc, "5A.2 Agent 执行策略", 2)
    add_table(
        doc,
        ["策略", "适用情况", "控制点"],
        [
            ["Function Calling", "模型支持结构化工具调用，适合生产默认策略。", "工具 Schema、参数校验、权限、最大迭代次数。"],
            ["ReAct", "需要显式推理轨迹或模型函数调用能力不足。", "思考-动作-观察记录、敏感信息脱敏、循环保护。"],
            ["Plan-and-Execute", "任务较长，需要先生成计划再分步执行。", "计划审批、步骤预算、失败回滚。"],
            ["Hybrid Workflow Agent", "工作流节点内部需要 Agent 自主完成子任务。", "节点边界、输入输出 Schema、超时和人工确认。"],
        ],
        [1.45, 2.35, 2.7],
    )
    add_heading(doc, "5A.3 用户体验", 2)
    for item in [
        "用户输入自然语言目标，例如：评估这份合同是否可以签署，并说明必须补充哪些材料。",
        "系统展示 Agent 计划：准备调用哪些本体、图谱、规则、技能和业务工具。",
        "运行中显示工具调用日志、证据、风险路径和成本/耗时。",
        "高影响动作暂停等待人工确认，如冻结付款、阻断签署、发起审批、对外发送通知。",
        "最终输出风险结论、决策建议、执行任务和可追溯推理轨迹。",
    ]:
        add_bullet(doc, item)


def section_ontology_linking(doc: Document) -> None:
    add_heading(doc, "6A. 本体链接机制", 1)
    add_para(doc, "无论是工作流节点还是 Agent 工具，都不能只是普通节点或普通函数；它们必须能链接智谱本体，才能沉淀成合同领域平台能力。")
    add_table(
        doc,
        ["链接对象", "链接位置", "作用"],
        [
            ["对象/属性", "输入节点、事实映射节点、抽取节点、工具参数。", "让事实知道自己属于 Contract.amount、Party.credit_rating 等字段。"],
            ["关系", "图谱查询节点、路径扩展工具、条件分支。", "限制查询路径，例如 Contract -> Party -> SealAuthorization。"],
            ["行为", "人工审批、执行任务、业务系统回写。", "把风险结论转成 RequireLegalReview、BlockSigning 等行为。"],
            ["规则", "规则执行节点、Agent rule.run 工具。", "记录规则命中、解释和版本。"],
            ["函数", "函数计算节点、function.evaluate 工具。", "复用智谱定义的高金额阈值、付款求和、最高风险分等函数。"],
            ["查询能力", "图谱节点、Agent graph 工具。", "复用智谱发布的 ontology_graph/local_graph 查询能力。"],
            ["发布包版本", "工作流版本、Agent 会话、语义案件。", "保证历史运行可复现。"],
        ],
        [1.25, 1.95, 3.3],
    )
    add_heading(doc, "6A.1 数据结构建议", 2)
    add_callout(
        doc,
        "ontology_refs",
        "{ package_id, package_version, refs: [{ ref_type: 'object|property|relation|action|rule|function|query_capability', code, id, required, usage: 'input|condition|tool|output' }] }",
        fill=LIGHT_BLUE,
    )


def section_updated_api(doc: Document) -> None:
    add_heading(doc, "7A. 双模式 API 增量", 1)
    add_table(
        doc,
        ["接口", "方法", "用途"],
        [
            ["/api/brain/workflows", "POST/GET", "创建和查询工作流定义。"],
            ["/api/brain/workflows/{id}", "GET/PUT", "查看和更新工作流画布、节点、边、本体引用。"],
            ["/api/brain/workflows/{id}/validate", "POST", "验证节点配置、变量连线、本体引用和权限。"],
            ["/api/brain/workflows/{id}/run", "POST", "运行工作流，生成 run_id 和节点轨迹。"],
            ["/api/brain/runs/{id}", "GET", "查看运行状态、节点输出、错误和轨迹。"],
            ["/api/brain/tools", "GET/POST", "注册 Agent 可调用工具，含本体工具、技能工具、业务工具。"],
            ["/api/brain/agents", "POST/GET", "创建 Agent 配置，绑定模型、工具集、系统提示词和权限。"],
            ["/api/brain/agents/{id}/chat", "POST", "以自然语言触发 Agent 自主工具调用。"],
            ["/api/brain/approvals", "GET/POST", "处理高影响工具调用或人工审批节点。"],
            ["/api/brain/ontology-refs/resolve", "POST", "校验并解析节点或工具上的本体引用。"],
        ],
        [2.1, 0.75, 3.65],
    )
    add_heading(doc, "7A.1 表结构增量", 2)
    add_table(
        doc,
        ["表", "关键字段", "说明"],
        [
            ["brain_workflows", "id, name, mode, package_id, version, status, canvas_json", "工作流定义，mode 支持 workflow/chatflow/hybrid。"],
            ["brain_workflow_nodes", "id, workflow_id, node_type, config, ontology_refs, position", "节点配置和本体链接。"],
            ["brain_workflow_edges", "id, workflow_id, source_node_id, target_node_id, condition", "节点连线和条件。"],
            ["brain_runs", "id, case_id, workflow_id, agent_id, mode, status, started_at, finished_at", "统一运行记录。"],
            ["brain_run_steps", "id, run_id, node_id, tool_id, input, output, status, trace", "节点执行或工具调用轨迹。"],
            ["brain_tools", "id, code, name, tool_type, schema, ontology_refs, permission_policy", "Agent 工具注册表。"],
            ["brain_agents", "id, name, model_config, tool_policy, system_prompt, max_iterations", "Agent 配置。"],
            ["brain_approvals", "id, run_id, step_id, action_type, payload, status, approved_by", "人工确认和审批。"],
        ],
        [1.45, 3.2, 1.85],
    )


def section_updated_mvp(doc: Document) -> None:
    add_heading(doc, "8A. 更新后的 MVP 范围", 1)
    add_table(
        doc,
        ["阶段", "目标", "必须交付"],
        [
            ["P1 工作流骨架", "能画流程、配置节点、绑定本体、运行简单 DAG。", "工作流 CRUD、节点/边 CRUD、节点测试、运行日志、ontology_refs 校验。"],
            ["P2 合同风控节点包", "覆盖合同风控常用节点。", "本体包加载、事实映射、图谱查询、规则执行、人工审批、执行任务、组织记忆节点。"],
            ["P3 Agent 工具模式", "能把本体和 skills 封装为工具，Agent 自主调用。", "工具注册表、Agent 配置、chat/run API、工具调用轨迹、迭代限制。"],
            ["P4 混合模式", "工作流节点可调用 Agent 子任务。", "Agent 节点、输入输出 Schema、超时/失败策略、人工确认。"],
            ["P5 运营治理", "可观测、可审计、可回放。", "运行详情、成本统计、审批、版本锁定、反馈回流智谱。"],
        ],
        [1.1, 1.65, 3.75],
    )
    add_callout(
        doc,
        "开发顺序建议",
        "先做工作流模式，因为它确定性强、便于测试；再做 Agent 工具模式，因为工具抽象可以复用工作流节点能力；最后做混合模式。",
        fill="F7FAFC",
    )


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)
    section_product_definition(doc)
    section_open_source_choice(doc)
    section_dual_mode_architecture(doc)
    section_workflow_mode(doc)
    doc.add_page_break()
    section_agent_mode(doc)
    section_ontology_linking(doc)
    section_scenarios(doc)
    section_user_experience(doc)
    doc.add_page_break()
    section_runtime_model(doc)
    section_reasoning(doc)
    doc.add_page_break()
    section_updated_api(doc)
    section_database(doc)
    doc.add_page_break()
    section_updated_mvp(doc)
    section_mvp_plan(doc)
    section_risks(doc)
    doc.add_page_break()
    add_heading(doc, "附录：参考资料", 1)
    add_table(
        doc,
        ["资料", "用于本设计的判断"],
        [
            ["Dify Docs - Workflow Nodes", "参考节点类型、画布输入、工具节点、Agent 节点、HTTP/Code/Template/Loop 等节点分类。"],
            ["Dify Docs - Agent Node", "参考 Agent 工具配置、Function Calling/ReAct、迭代限制、工具输出和推理轨迹。"],
            ["LangGraph Docs - Workflows and Agents", "参考 workflow 和 agent 的边界：workflow 是预定义路径，agent 是动态工具使用。"],
            ["LangGraph Docs - Overview / Graph API", "参考 StateGraph、节点、边、状态、持久执行、人机协同和调试能力。"],
        ],
        [2.2, 4.3],
    )
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)

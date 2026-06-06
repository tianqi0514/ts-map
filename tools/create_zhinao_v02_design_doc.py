from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

sys.path.append(str(Path(__file__).parent))
from create_zhinao_design_doc import (  # noqa: E402
    DARK_BLUE,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    RED,
    add_bullet,
    add_callout,
    add_heading,
    add_number,
    add_para,
    add_table,
    configure_doc,
    paragraph_border_bottom,
    set_para_spacing,
    set_run_font,
)


OUT = Path("/Users/tianqi/Documents/CODEX-MAP/docs/传神智脑设计/传神智脑-合同语义网络决策平台PRD与详细设计-v0.2.docx")


def add_title_page(doc: Document) -> None:
    add_para(doc, "PRD 与详细设计 · v0.2", size=11, color=MUTED, bold=True, after=8)
    title = doc.add_paragraph()
    set_para_spacing(title, after=4)
    run = title.add_run("传神智脑：合同语义网络决策平台")
    set_run_font(run, size=24, color=INK, bold=True)
    subtitle = add_para(
        doc,
        "面向组织层面的合同推理、决策、风控与执行协同；承接智谱发布包，不做单员工替身型 Agent",
        size=13,
        color=MUTED,
        after=16,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for label, value in [
        ("适用阶段", "智谱本体管理完成后，智脑一期产品设计与开发拆解"),
        ("业务域", "合同风控、重大合同决策、主体风险穿透、履约与付款协同"),
        ("核心依赖", "智谱发布包、Neo4j 语义网络、PostgreSQL 运行态数据、业务系统事实源"),
        ("核心输出", "风险判断、决策建议、执行计划、推理轨迹、组织记忆"),
        ("版本日期", "2026-06-07"),
    ]:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=11, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=11, color=INK)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)
    add_callout(
        doc,
        "核心判断",
        "智脑的价值不在于替某个员工执行单点任务，而在于把合同事实、组织关系、规则、外部风险和历史决策连接成可推理的语义网络，让企业能回答“这笔交易能不能做、风险如何传导、该由谁决策、后续动作是否闭环”。",
        fill="EEF5FF",
    )


def section_product_definition(doc: Document) -> None:
    add_heading(doc, "1. 产品定位与边界", 1)
    add_para(
        doc,
        "传神智脑是合同领域的组织级语义网络决策平台。它以智谱发布的本体包为语义底座，将真实合同、主体、条款、义务、付款、审批、附件、外部风险与历史处置结果组织成运行时语义网络，并在网络上进行推理、风控、决策和执行编排。",
    )
    add_table(
        doc,
        ["维度", "不做什么", "要做什么"],
        [
            ["产品价值", "不做单份合同问答助手或个人写作助手。", "做组织级合同风险识别、决策建议、执行闭环和经营洞察。"],
            ["Agent 定位", "不把 Agent 理解为替员工点按钮的自动化脚本。", "把 Agent 理解为基于语义网络协同推理、调度工具、记录证据的决策单元。"],
            ["系统边界", "不替代 CLM、ERP、OA、CRM 成为业务主库。", "连接事实源，做语义对齐、风险判断、行动建议和结果回写。"],
            ["规则治理", "不把规则硬编码进流程，避免越做越僵。", "承接智谱中的行为与规则，用运行时事实、图路径和证据置信度解释规则命中。"],
            ["自动化边界", "不默认自动执行高影响动作。", "高风险、阻断、付款、签署等动作必须保留人工确认和审计链。"],
        ],
        [1.15, 2.7, 2.65],
    )
    add_callout(
        doc,
        "一期设计原则",
        "体验上要简单：用户看到的是合同、风险、决策、动作和证据；复杂的图推理、规则编排、版本匹配和置信度计算都在后台完成。",
        fill="F7FAFC",
    )


def section_scenarios(doc: Document) -> None:
    add_heading(doc, "2. 一期业务场景", 1)
    add_para(doc, "智脑一期不做泛化 Agent 平台，建议聚焦合同风控中能体现组织级价值的 4 个场景。")
    add_table(
        doc,
        ["场景", "业务问题", "核心输入", "核心输出"],
        [
            ["重大合同决策", "金额大、条款复杂、责任重大时，合同是否可签？需要谁审批？", "合同事实、金额、条款、主体、付款计划、智谱规则。", "风险等级、阻断点、审批路径、条款修改建议。"],
            ["主体风险穿透", "对手方背后是否有关联、失信、资质过期或历史履约问题？", "主体信息、股权关系、司法/工商/舆情、历史合同。", "主体风险图、证据链、交易限制建议。"],
            ["履约与付款风控", "付款条件、交付里程碑、验收、发票是否一致？", "义务、付款计划、交付物、验收记录、发票记录。", "异常路径、冻结/补证建议、责任部门。"],
            ["政策变更影响", "新制度、新法规、新监管口径会影响哪些存量合同？", "政策条款、合同条款、业务线、主体、风险规则。", "影响合同清单、整改计划、优先级。"],
        ],
        [1.25, 1.85, 1.75, 1.65],
    )
    add_heading(doc, "2.1 MVP 首选闭环", 2)
    for item in [
        "首选闭环：重大合同决策 + 主体风险穿透。它最能体现组织层面的风险传导、审批决策和证据解释。",
        "演示数据可沿用当前智谱的“合同风控决策本体”：合同、合同主体、条款、付款计划、风险项、审批记录、附件证据、法规政策、签署授权。",
        "一期先支持结构化事实录入和样例数据导入，再逐步接入合同文本抽取和业务系统连接器。",
    ]:
        add_bullet(doc, item)


def section_user_experience(doc: Document) -> None:
    add_heading(doc, "3. 信息架构与页面设计", 1)
    add_para(doc, "左侧一级导航新增 `智脑`。智脑下的二级页面不暴露图数据库、Agent 编排等技术概念，而以合同决策工作流组织。")
    add_table(
        doc,
        ["二级页面", "用户任务", "核心组件", "一期 CRUD/操作"],
        [
            ["工作台", "看到待决策事项、重大风险、执行卡点。", "风险概览、待办队列、重大合同卡片、今日超期动作。", "查询、筛选、进入案件、标记关注。"],
            ["语义案件", "管理一次合同推理/决策任务。", "案件列表、案件详情、输入事实、决策结果、推理轨迹。", "新增、编辑、关闭、复制、归档。"],
            ["合同语义网络", "查看合同背后的主体、条款、义务、付款和风险路径。", "图谱视图、路径面板、证据面板、对象详情。", "查询、展开、固定节点、导出路径。"],
            ["场景推理舱", "选择场景并运行推理。", "场景选择、输入事实、激活语义包、运行记录。", "创建运行、重跑、保存草稿、查看结果。"],
            ["风险与决策", "处理风险发现和决策建议。", "风险列表、决策卡、证据链、人工确认。", "确认、驳回、升级、转执行。"],
            ["执行协同", "跟踪建议动作是否闭环。", "行动清单、责任人、截止时间、外部系统状态。", "创建、分派、催办、关闭。"],
            ["组织记忆", "复盘规则效果和人工反馈。", "案例库、反馈记录、规则效果、智谱优化建议。", "新增案例、采纳反馈、生成回流建议。"],
        ],
        [1.15, 1.55, 2.05, 1.75],
    )
    add_heading(doc, "3.1 核心页面细节", 2)
    add_table(
        doc,
        ["页面", "布局建议", "关键交互"],
        [
            ["语义案件详情", "左侧案件摘要，中间风险/决策/证据，右侧推理轨迹。", "用户先看结论，再展开证据和图路径；支持人工确认和驳回。"],
            ["合同语义网络", "浅色图谱为主，右侧详情抽屉，底部显示路径解释。", "点击合同/主体/条款节点展示事实；点击风险路径展示规则和证据。"],
            ["场景推理舱", "上方场景与语义包版本，中间输入事实，下方运行结果。", "用户不配置流程，只选择场景并补齐必要事实。"],
            ["风险与决策", "按风险等级、业务线、责任部门、状态组织表格。", "高风险项需要人工确认；所有结论必须可追溯。"],
            ["执行协同", "以任务表为主，强调责任人、状态、外部系统回写。", "阻断、付款冻结、补证、审批等动作必须有状态闭环。"],
        ],
        [1.35, 2.35, 2.8],
    )


def section_runtime_model(doc: Document) -> None:
    add_heading(doc, "4. 运行态领域模型", 1)
    add_para(doc, "智谱管理的是本体定义，智脑管理的是运行态事实、推理案件、决策结果和执行闭环。两者通过发布包版本关联。")
    add_table(
        doc,
        ["模型", "用途", "关键字段"],
        [
            ["SemanticPackage", "智谱发布包快照，智脑运行时只读引用。", "space_id、package_id、version、objects、relations、actions、rules、functions、published_at。"],
            ["RuntimeFact", "合同事实、主体事实、付款事实等运行时数据。", "case_id、ontology_type、ontology_code、value、source、confidence、evidence_id。"],
            ["SemanticCase", "一次合同决策/推理任务。", "case_no、scenario_type、space_id、package_id、subject_ref、status、risk_level。"],
            ["Evidence", "支持判断的证据。", "source_type、source_uri、quote、structured_value、confidence、captured_at。"],
            ["RiskFinding", "风险发现。", "case_id、risk_type、severity、score、rule_id、path_refs、evidence_refs、recommendation。"],
            ["DecisionResult", "决策建议或人工确认后的结果。", "decision、reason、recommended_actions、requires_review、confirmed_by、confirmed_at。"],
            ["ReasoningTrace", "推理过程。", "step_no、step_type、input_refs、operation、output、confidence、elapsed_ms。"],
            ["ExecutionTask", "执行动作。", "action_code、owner_role、owner_user、due_at、status、external_system、callback_payload。"],
            ["FeedbackRecord", "复盘反馈。", "accepted、override_reason、final_outcome、ontology_feedback、created_by。"],
        ],
        [1.3, 2.05, 3.15],
    )
    add_heading(doc, "4.1 Neo4j 图模型", 2)
    add_para(doc, "智脑需要在 Neo4j 中维护运行态语义图。PostgreSQL 保存业务记录和审计，Neo4j 负责多跳查询、路径解释和影响面计算。")
    add_table(
        doc,
        ["节点/边", "类型", "说明"],
        [
            ["(:Case)", "节点", "语义案件，关联场景、合同和决策。"],
            ["(:Fact)", "节点", "运行态事实，可映射到智谱对象或属性。"],
            ["(:Evidence)", "节点", "证据，来源于合同文本、系统数据或外部数据。"],
            ["(:Risk)", "节点", "风险发现，连接规则、证据和对象路径。"],
            ["(:Decision)", "节点", "决策结果。"],
            ["(:Task)", "节点", "执行动作。"],
            ["(:Fact)-[:INSTANCE_OF]->(:OntologyNode)", "边", "运行事实实例化智谱本体对象。"],
            ["(:Risk)-[:SUPPORTED_BY]->(:Evidence)", "边", "风险由证据支撑。"],
            ["(:Risk)-[:TRIGGERED_BY]->(:OntologyNode)", "边", "风险由智谱规则或行为触发。"],
            ["(:Decision)-[:GENERATES]->(:Task)", "边", "决策生成执行动作。"],
        ],
        [2.0, 0.85, 3.65],
    )


def section_reasoning(doc: Document) -> None:
    add_heading(doc, "5. 推理与决策链路", 1)
    add_number(doc, "加载语义包：选择智谱发布包版本，锁定本次推理使用的对象、关系、函数、行为和规则。")
    add_number(doc, "事实接入：从合同文本、台账、审批、付款、发票、主体和外部风险数据生成 RuntimeFact。")
    add_number(doc, "语义对齐：将 RuntimeFact 映射到本体对象、属性和关系，形成运行态语义网络。")
    add_number(doc, "场景识别：根据输入事实和用户选择激活重大合同、主体穿透、履约付款或政策影响场景。")
    add_number(doc, "图扩展：围绕合同、主体、条款、义务、付款、风险做多跳遍历，收集上下文和证据。")
    add_number(doc, "函数计算：调用智谱发布的函数，如高金额阈值、付款计划求和、最高风险分等。")
    add_number(doc, "规则推理：执行硬规则、软规则、冲突处理和置信度计算，生成 RiskFinding。")
    add_number(doc, "决策生成：根据风险、证据、组织授权和历史反馈生成 DecisionResult 和 ExecutionTask。")
    add_number(doc, "人工确认：高影响动作必须由责任人确认，确认结果进入审计和组织记忆。")
    add_number(doc, "回流智谱：发现新对象、关系、属性、规则或规则效果问题时，生成智谱优化建议。")
    add_heading(doc, "5.1 冲突与置信度", 2)
    add_table(
        doc,
        ["类型", "处理方式"],
        [
            ["事实冲突", "同一字段来自多个来源且值不一致时，按来源可信度、时间新鲜度和人工确认状态排序。"],
            ["规则冲突", "硬规则优先于软规则；阻断类规则优先于建议类规则；冲突必须展示给用户。"],
            ["证据不足", "不得强行给确定结论，应输出“需补充材料”动作和缺失证据清单。"],
            ["低置信度抽取", "参与推理但降低结论置信度，高风险场景要求人工确认。"],
            ["人工覆盖", "允许人工覆盖智脑建议，但必须填写原因，进入组织记忆用于复盘。"],
        ],
        [1.45, 5.05],
    )


def section_api(doc: Document) -> None:
    add_heading(doc, "6. 后端 API 设计", 1)
    add_para(doc, "一期 API 以语义案件为中心，保持简单可测。每个接口都应有 pytest 覆盖，并在本地演示数据上可直接运行。")
    add_table(
        doc,
        ["接口", "方法", "用途", "测试要点"],
        [
            ["/api/brain/packages", "GET", "列出可用于智脑的智谱发布包。", "只返回验证通过并发布的包。"],
            ["/api/brain/cases", "POST/GET", "创建和查询语义案件。", "创建时必须绑定 space_id/package_id/scenario_type。"],
            ["/api/brain/cases/{id}", "GET/PUT", "查看和更新案件。", "状态流转、字段更新、审计记录。"],
            ["/api/brain/cases/{id}/facts", "POST/GET", "录入或查询运行态事实。", "事实能映射到本体 code，置信度合法。"],
            ["/api/brain/cases/{id}/run", "POST", "运行场景推理。", "生成风险、决策、轨迹；重复运行可追踪。"],
            ["/api/brain/cases/{id}/graph", "GET", "读取案件语义网络。", "返回节点、边、风险路径、证据路径。"],
            ["/api/brain/risks", "GET", "风险发现列表。", "支持场景、等级、责任部门、状态筛选。"],
            ["/api/brain/decisions/{id}/confirm", "POST", "人工确认或驳回决策。", "高影响动作必须确认；保存覆盖原因。"],
            ["/api/brain/tasks", "GET/POST", "执行动作管理。", "创建、分派、催办、关闭。"],
            ["/api/brain/feedback", "POST", "记录人工反馈和回流建议。", "能生成智谱待优化项。"],
        ],
        [1.85, 0.65, 2.1, 1.9],
    )
    add_heading(doc, "6.1 推理接口返回示例", 2)
    add_callout(
        doc,
        "POST /api/brain/cases/{id}/run 返回结构",
        "{ case_id, scenario_type, package_id, risk_level, risks: [{risk_type, severity, score, evidence_refs, path_refs}], decision: {decision, reason, recommended_actions}, traces: [{step_no, step_type, output, confidence}], tasks: [{action_code, owner_role, due_at, status}] }",
        fill=LIGHT_BLUE,
    )


def section_database(doc: Document) -> None:
    add_heading(doc, "7. PostgreSQL 表设计建议", 1)
    add_table(
        doc,
        ["表", "关键字段", "说明"],
        [
            ["brain_semantic_packages", "id, space_id, package_id, version, snapshot, status, published_at", "智谱发布包快照，智脑运行时只读。"],
            ["brain_semantic_cases", "id, case_no, space_id, package_id, scenario_type, subject_type, subject_id, status, risk_level", "语义案件主表。"],
            ["brain_runtime_facts", "id, case_id, ontology_type, ontology_code, field_code, value, confidence, evidence_id", "运行态事实。"],
            ["brain_evidences", "id, case_id, source_type, source_uri, quote, structured_value, confidence", "证据表。"],
            ["brain_risk_findings", "id, case_id, risk_type, severity, score, rule_code, recommendation, status", "风险发现。"],
            ["brain_decision_results", "id, case_id, decision, reason, recommended_actions, confirmed_by, confirmed_at", "决策结果。"],
            ["brain_reasoning_traces", "id, case_id, run_id, step_no, step_type, input_refs, output, confidence", "推理轨迹。"],
            ["brain_execution_tasks", "id, case_id, action_code, owner_role, owner_user, due_at, status, callback_payload", "执行动作。"],
            ["brain_feedback_records", "id, case_id, target_type, target_id, accepted, override_reason, ontology_feedback", "复盘反馈。"],
        ],
        [1.55, 3.15, 1.8],
    )
    add_callout(
        doc,
        "数据一致性",
        "智脑运行态数据必须记录 package_id/version。即使智谱后续发布新版本，历史案件仍能复现当时的推理依据。",
        fill="FFF4E5",
    )


def section_mvp_plan(doc: Document) -> None:
    add_heading(doc, "8. MVP 开发拆解", 1)
    add_table(
        doc,
        ["阶段", "目标", "后端交付", "前端交付"],
        [
            ["P1 语义案件", "能创建合同风控案件并录入事实。", "packages/cases/facts API，基础表结构和测试。", "智脑导航、工作台、语义案件列表与详情。"],
            ["P2 网络与推理", "能基于智谱包生成风险和决策。", "run/graph API，规则执行、函数调用、ReasoningTrace。", "场景推理舱、推理结果、证据链、风险路径。"],
            ["P3 执行闭环", "能把建议动作转成任务并跟踪状态。", "tasks/confirm/feedback API，审计和回写模拟。", "风险与决策台、执行协同台。"],
            ["P4 组织记忆", "能复盘人工反馈并生成智谱优化建议。", "feedback analytics，ontology_feedback 队列。", "组织记忆、规则效果分析、回流建议列表。"],
        ],
        [0.95, 1.45, 2.25, 1.85],
    )
    add_heading(doc, "8.1 一期验收标准", 2)
    for item in [
        "从智谱发布包读取合同风控本体，并创建一条重大合同决策案件。",
        "录入合同金额、主体信用、签署授权、付款计划、附件核验等事实。",
        "运行推理后生成至少 3 类风险：高金额审批、低信用主体、缺失签署授权或附件未核验。",
        "每个风险能展示证据、规则、图路径和推理步骤。",
        "高风险决策能转成执行任务，支持确认、驳回、关闭。",
        "人工反馈能进入组织记忆，并生成智谱规则/属性优化建议。",
        "所有新增 API 都有自动化测试，演示数据可一键初始化。",
    ]:
        add_bullet(doc, item)


def section_risks(doc: Document) -> None:
    add_heading(doc, "9. 风险与约束", 1)
    add_table(
        doc,
        ["风险", "表现", "控制策略"],
        [
            ["数据质量不足", "合同事实缺失、抽取错误、外部数据滞后。", "证据置信度、缺失事实提示、人工确认和来源优先级。"],
            ["推理不可解释", "用户只看到结论，无法信任系统。", "强制返回规则、证据、图路径和推理轨迹。"],
            ["自动化越权", "系统自动执行高影响动作。", "动作分级，高影响动作必须人工确认。"],
            ["智谱版本漂移", "历史案件无法复现。", "案件绑定 package_id/version，发布包快照只读。"],
            ["场景膨胀", "过早做成通用 Agent 平台，主线发散。", "一期锁定重大合同决策和主体风险穿透。"],
        ],
        [1.25, 2.25, 3.0],
    )


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)
    section_product_definition(doc)
    section_scenarios(doc)
    doc.add_page_break()
    section_user_experience(doc)
    section_runtime_model(doc)
    section_reasoning(doc)
    section_api(doc)
    section_database(doc)
    doc.add_page_break()
    section_mvp_plan(doc)
    section_risks(doc)
    add_heading(doc, "附录：和智谱模块的分工", 1)
    add_table(
        doc,
        ["能力", "智谱负责", "智脑负责"],
        [
            ["对象/关系/行为/规则", "定义、维护、验证、发布。", "按发布包版本只读消费。"],
            ["属性与规则归属", "属性归对象，规则归行为。", "将运行时事实映射到对象属性，将规则命中转为风险和动作。"],
            ["场景", "不在智谱 MVP 中定义。", "在智脑中作为组织决策任务定义和运行。"],
            ["图谱", "展示本体对象与关系。", "展示运行时事实、风险、证据、决策和执行路径。"],
            ["版本", "发布可调用语义包。", "案件绑定语义包版本，支持历史复现。"],
        ],
        [1.4, 2.55, 2.55],
    )
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
